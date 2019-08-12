import sys, os, docx
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QMessageBox
from req_ui import *
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = {}
personsArr = []
input = []

def makeDict():
    global d, personsArr, input
    with open("bin" + os.sep + "input.txt", encoding = "utf-8") as f:
        input = f.read().split("\n")

    d = {}
    personsArr = []

    d["ФИО"] = []
    d["Адрес"] = []

    for line in input:
        if line != "":
            title = line.split(": ",1)[0]
            if "Файл" in title and title != "Файл ИЦ":
                d["Файл"] = line.split(": ",1)[1]
            if title == "ФИО" or title == "Адрес":
                d[title].append(line.split(": ",1)[1])
            else:
                d[title] = line.split(": ",1)[1]

    # создаём массив лиц
    for i in range(len(d["ФИО"])):
        per = ""
        per += "ФИО: " + d["ФИО"][i] + "\n"
        per += "Адрес: " + d["Адрес"][i] + "\n"
        personsArr.append(per)


class MyWin(QtWidgets.QMainWindow):
    def __init__ (self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        # считываем файл input.txt
        try:
            makeDict()
        except Exception as e:
            QMessageBox.warning(myapp, 'Ошибка', "Error: " + str(e))

        # Заполняем поля сохраненными значениями
        self.fillAllFields()

        # Обработчик событий
        self.ui.mpRadioButton.clicked.connect(self.changeFabula)
        self.ui.udRadioButton.clicked.connect(self.changeFabula)
        self.ui.readyButton.clicked.connect(self.reqFunc)
        self.ui.addPersonButton.clicked.connect(self.addPerson)
        self.ui.delPersonButton.clicked.connect(self.delPerson)
        self.ui.clearPersonButton.clicked.connect(self.clearPerson)
        self.ui.defaultButton.clicked.connect(self.defaultAll)

    def writePersonInput(self):
        t = personsArr[:]
        out = ""
        count = 0
        for el in t:
            count += 1
            el = el.replace("ФИО: ", "").split("\nАдрес:")
            el = str(el[0] + "," + el[1])
            out += str(count) + ") " + el + "\n"
        self.ui.personInput.setPlainText(out)

    def fillAllFields(self):
        # Заполняем поля сохраненными значениями
        self.ui.dataEdit.setText(d["Дата"])
        self.ui.numberEdit.setText(d["Номер материала/дела"])
        self.ui.placeEdit.setText(d["Населенный пункт"])
        if int(d["Материал/Дело[0/1]"]) == 0:
            self.ui.mpRadioButton.setChecked(True)
        if int(d["Материал/Дело[0/1]"]) == 1:
            self.ui.udRadioButton.setChecked(True)
        self.changeFabula()

        # выводим массив лиц в поле
        self.writePersonInput()

        # данные следователя (знаки *-* нужны для корректного отображения считывания строки)
        self.ui.positionInput.setPlainText(d["Должность"].replace("*-*", "\n"))
        self.ui.rankEdit.setText(d["Звание"])
        self.ui.sledNameEdit.setText(d["Имя следователя"])

    def defaultAll(self):
        try:
            with open("bin" + os.sep + "default.txt", encoding = "utf-8") as f:
                t = f.read()
            with open("bin" + os.sep + "input.txt", "w", encoding = "utf-8") as f:
                f.write(t)
            makeDict(); # заново формируем словарь и массив лиц
        except Exception as e:
            QMessageBox.warning(myapp, 'Ошибка', "Error: " + str(e))
        self.fillAllFields()

    def clearPerson(self):
        # очищает поле и массив лиц
        global personsArr
        personsArr = []
        self.ui.personInput.setPlainText("")

    def delPerson(self):
        # удаляем последнее лицо из массива лиц и выводим в поле
        try:
            personsArr.pop()
        except Exception:
            print("personsArr is empty")
        self.writePersonInput()

    def addPerson(self):
        # добавляет лицо в массив и выводит в специальное поле
        per = ""
        per += "ФИО: " + self.ui.fioEdit.text()
        per += ", " + self.ui.bornEdit.text() + " г.р.\n"
        per += "Адрес: проживающий(-ая) по адресу: " + self.ui.adressTextEdit.toPlainText() + ";\n"
        personsArr.append(per)
        self.writePersonInput()

    def changeFabula(self):
        if self.ui.mpRadioButton.isChecked():
            fab = "Фабула по материалу"
        if self.ui.udRadioButton.isChecked():
            fab = "Фабула по делу"
        self.ui.fabulaInput.setPlainText(d[fab])

    def reqFunc(self):
        try:
            self.writeFunc()
            makeDict()
            doc = docx.Document("bin" + os.sep + d["Файл"] + ".docx")

            # работа с таблицами (нужны поля Date и Country)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text == "DATE": # номер дела или материала
                                    run.text = d["Дата"] + " " * 16 + d["Номер материала/дела"]

                                if "COUNTRY" in run.text: # населенный пункт
                                    run.text = run.text[:-len("COUNTRY")-1] + d["Населенный пункт"] + "»"

            for p in doc.paragraphs:
                for run in p.runs:
                    if run.text == "DATE": # номер дела или материала
                        run.text = d["Дата"] + " " * 16 + d["Номер материала/дела"]

                    if "COUNTRY" in run.text: # населенный пункт
                        run.text = run.text[:-len("COUNTRY")-1] + d["Населенный пункт"] + "»"

                    if int(d["Материал/Дело[0/1]"]) == 0:
                        fab = "Фабула по материалу"
                    else:
                        fab = "Фабула по делу"

                    if "FABULA" in run.text:
                        run.text = d[fab] + run.text[len("FABULA"):]

                    if run.text == "FIO": # имя
                        i = 0
                        while i < len(d["ФИО"]): # добавляем параграфы
                            i += 1
                            if i == len(d["ФИО"]):
                                newParagraph = p
                            else:
                                newParagraph = p.insert_paragraph_before()
                            newParagraph.add_run()
                            newParagraph.runs[0].text = "\t" + str(i//2 + 1) + ") " + d["ФИО"][i//2]
                            newParagraph.runs[0].bold = True
                            newParagraph.runs[0].font.name = "Times New Roman"
                            newParagraph.runs[0].font.size = Pt(14)
                            newParagraph.add_run()
                            newParagraph.runs[1].text = ", " + d["Адрес"][i//2]
                            newParagraph.runs[1].font.name = "Times New Roman"
                            newParagraph.runs[1].font.size = Pt(14)
                            newParagraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    if "POSITION" in run.text:
                        run.text = d["Должность"].replace("*-*", "\n")
                    if "RANK" in run.text:
                        run.text = d["Звание"]
                    if "NAME" in run.text:
                        run.text = d["Имя следователя"]

            doc.save('Запросы готовые.docx')

            # требования ИЦ, ГИАЦ
            for count in range(len(d["ФИО"])):
                doc = docx.Document("bin" + os.sep + d["Файл ИЦ"] + ".docx") # файл ИЦ
                for p in doc.paragraphs:
                    for run in p.runs:
                        if "FAMILIA" in run.text:
                            run.text = d["ФИО"][count].split()[0]
                        if "NAMEMIDDLE" in run.text:
                            run.text = d["ФИО"][count].split()[1] + " " + d["ФИО"][count].split()[2][:-1]
                        if "YEAR" in run.text:
                            run.text = d["ФИО"][count].split(",")[1]
                        if "ADRESS" in run.text:
                            run.text =  d["Адрес"][count].split(":")[1]
                        if "NUMBER" in run.text:
                            run.text =  "№" + d["Номер материала/дела"]
                        if "CURRENT" in run.text:
                            run.text =  " " * 6 + d["Дата"]
                        if "POSITION" in run.text:
                            run.text = d["Должность"].replace("*-*", "\n")
                        if "RANK" in run.text:
                            run.text = d["Звание"]
                        if "NAME" in run.text:
                            run.text = d["Имя следователя"]

                doc.save('Требование ИЦ, ГИАЦ ' + str(count + 1) + " " + d["ФИО"][count].split()[0] + '.docx')

            QMessageBox.information(myapp, "Done", "Файлы созданы успешно")
        except Exception as e:
            # всплывающее окно с ошибкой
            QMessageBox.warning(myapp, 'Ошибка', "Error: " + str(e))

    def writeFunc(self):
        # Присваиваем значения полей в переменную input, котрую затем запишем в файл
        for line in input:
            if line != "":
                title = line.split(": ",1)[0]
                if title == "ФИО" or title == "Адрес":
                    input.remove(line)
                if title == "Дата":
                    input[input.index(line)] = title + ": " + self.ui.dataEdit.text()
                if title == "Материал/Дело[0/1]":
                    if self.ui.mpRadioButton.isChecked():
                        input[input.index(line)] = title + ": " + "0"
                    else:
                        input[input.index(line)] = title + ": " + "1"
                if title == "Номер материала/дела":
                    input[input.index(line)] = title + ": " + self.ui.numberEdit.text()

                if title == "Фабула по материалу" and self.ui.mpRadioButton.isChecked():
                    input[input.index(line)] = title + ": " + self.ui.fabulaInput.toPlainText()
                if title == "Фабула по делу" and self.ui.udRadioButton.isChecked():
                    input[input.index(line)] = title + ": " + self.ui.fabulaInput.toPlainText()

                if title == "Населенный пункт":
                    input[input.index(line)] = title + ": " + self.ui.placeEdit.text()

                if title == "Должность":
                	input[input.index(line)] = title + ": " + self.ui.positionInput.toPlainText().replace("\n", "*-*")

                if title == "Звание":
                	input[input.index(line)] = title + ": " + self.ui.rankEdit.text()

                if title == "Имя следователя":
                	input[input.index(line)] = title + ": " + self.ui.sledNameEdit.text()

        # очищаем input от пустых элементов
        while "" in input:
            input.remove("")

        # вписываем в конец input содержимое поля personInput
        for person in personsArr:
            input.append(person.replace("Адрес", "\nАдрес"))

        # добавляем в input пустые элементы через каждую строку, чтобы корректно срабатывало последующее считывание из файла
        tInput = input
        for i in range(len(tInput) * 2 - 1):
            if i % 2 != 0:
                input.insert(i, "")

        # отредактированную переменную записываем в файл
        with open("bin" + os.sep + "input.txt", "w", encoding = "utf-8") as f:
            f.write("\n".join(input))

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    myapp = MyWin()
    myapp.show()
    myapp.setFixedSize(myapp.geometry().width(),myapp.geometry().height());
    sys.exit(app.exec_())
